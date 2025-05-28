import os
import subprocess
from datetime import datetime
from io import BytesIO
import yaml
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


file_input ="src/bruno.bronosky.resume.yaml"
file_output ="Bruno_Bronosky_Resume_From_YAML.docx"
section_breaks = {}

# Load YAML data
with open(file_input, "r") as f:
    docs = yaml.safe_load_all(f)
    for doc in docs:
        if doc and "$schema" in doc:
            resume = doc
            break

if resume is None:
    raise ValueError("Resume YAML document with $schema not found")

doc = Document()

def datefmt(date_str):
    return date_str if date_str.lower() == "present" else datetime.strptime(date_str, "%Y-%m-%d").strftime("%Y-%m")

def get_page_count(doc):
    doc.save(file_output)
    pdf_path = os.path.splitext(file_output)[0] + ".pdf"
    # Convert .docx to .pdf using LibreOffice
    print(f"ONO before")
    subprocess.call([
        "libreoffice", "--headless", "--convert-to", "pdf", file_output,
        "--outdir", os.path.dirname(file_output) or "."
    ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    print(f"ONO after")
    # Extract page count using pdfinfo
    pdfinfo = subprocess.Popen(["pdfinfo", pdf_path], stdout=subprocess.PIPE, stderr=subprocess.DEVNULL)
    awk = subprocess.Popen(["awk", "/Pages:/ {print $2}"], stdin=pdfinfo.stdout, stdout=subprocess.PIPE, text=True)

    pdfinfo.stdout.close()  # Allow pdfinfo to receive SIGPIPE if awk exits
    page_count_str, _ = awk.communicate()
    page_count = int(page_count_str.strip())
    return page_count

# Set margins for compact layout
section = doc.sections[0]
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

# Font defaults
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(9)
doc.styles['Heading 1'].font.size = Pt(12)
doc.styles['Heading 2'].font.size = Pt(10)
#breakpoint()

if "MySectionStyle" not in doc.styles:
    new_style = doc.styles.add_style("MySectionStyle", WD_STYLE_TYPE.PARAGRAPH)
    base = doc.styles["Normal"]
    new_style.base_style = base
    new_style.paragraph_format.space_after = Pt(0)

if "MyBulletStyle" not in doc.styles:
    new_style = doc.styles.add_style("MyBulletStyle", WD_STYLE_TYPE.PARAGRAPH)
    base = doc.styles["List Bullet"]
    new_style.base_style = base
    new_style.paragraph_format.space_after = Pt(10)

if "CustomLabel" not in doc.styles:
    font = doc.styles.add_style("CustomLabel", WD_STYLE_TYPE.CHARACTER).font
    #font.name = 'Arial'
    #font.size = Pt(12)
    font.bold = True
    #font.color.rgb = RGBColor(0x36, 0x5F, 0x91)  # Blue
    font.color.rgb = doc.styles['Title'].font.color.rgb

# Header
doc.add_heading(resume["basics"]["name"], level=0)
p = doc.add_paragraph()
p.add_run(f"{resume['basics']['label']}\n")
p.add_run(f"{resume['basics']['location']['city']}, {resume['basics']['location']['region']} {resume['basics']['location']['postalCode']} | {resume['basics']['phone']} | {resume['basics']['email']}")

# Hyperlinks (manually, or just as plain text for now)
p = doc.add_paragraph()
for idx, profile in enumerate(resume["basics"]["profiles"]):
    is_last = idx == len(resume["basics"]["profiles"]) - 1
    delimiter = "" if is_last else "\n"
    p.add_run(f"{profile['network']}: {profile['url']}{delimiter}")

# Summary
doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(resume["basics"]["summary"].strip())

# Skills
doc.add_heading("Skills", level=1)
p = doc.add_paragraph()

for idx, skill in enumerate(resume.get("skills", [])):
    is_last = idx == len(resume.get("skills", [])) - 1
    delimiter = "" if is_last else "\n"
    keywords = ", ".join(skill.get("keywords", []))
    p.add_run(f"{skill['name']}").style = doc.styles["CustomLabel"]
    p.add_run(f": {keywords}{delimiter}")

# Experience
def copy_doc(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    doc_copy = Document(buffer)
    return doc_copy

def add_experience(doc, job, previous_experience):
    d = copy_doc(doc)
    pages_before = get_page_count(d)
    add_experience_raw(d, job)
    pages_after = get_page_count(d)
    did_break = False
    if pages_after > pages_before:
        print(f"Page Break ({pages_after} pages) in {job['position']} {job['name']}")
        #doc.add_page_break()
        add_page_break_after_paragraph(previous_experience)
        did_break = True
    else:
        print(f"Unbroken ({pages_after} pages) in {job['position']} {job['name']}")
    p = add_experience_raw(doc, job)
    return [p, did_break]

def add_experience_raw(doc, job):
    title = job["position"]
    company = job["name"]
    dates = f"{datefmt(job.get('startDate', ''))} – {datefmt(job.get('endDate', ''))}"
    p = doc.add_paragraph(style="MySectionStyle")
    p.add_run(f"{company} | ")
    p.add_run(f"{title}").style = doc.styles["CustomLabel"]
    p.add_run(f" | {dates}")
    for line in job.get("highlights", []):
        p = doc.add_paragraph(f"{line}", style="MyBulletStyle")
    return_paragraph = p
    return return_paragraph

def add_page_break_after_paragraph(p):
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

doc.add_heading("Professional Experience", level=1)
previous_experience = []
for idx, job in enumerate(resume["work"]):
    if job.get("display", "") != "None":
        previous_experience, did_break = add_experience(doc, job, previous_experience)
        breaks = section_breaks.get('experience', [])
        breaks.append(idx)
        section_breaks['experience'] = breaks

# Education
#doc.add_heading("Education", level=1)
#for edu in resume.get("education", []):
#    edu_line = f"{edu['institution']}, {edu.get('startDate', '')} – {edu.get('endDate', '')}"
#    doc.add_paragraph(f"{edu['area']} — {edu_line}")

def add_markdown_paragraph(doc, md_text):
    from markdown import markdown
    from html4docx import HtmlToDocx
    from docx import Document
    html = markdown(md_text)
    parser = HtmlToDocx()
    parser.add_html_to_document(html, doc)

#md_text = "**Bold**, *italic*, and a [link](https://example.com)."
#add_markdown_paragraph(doc, md_text)

# Save
doc.save(file_output)
print(f'Successfully generated {file_output}')
