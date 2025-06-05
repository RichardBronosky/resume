"""
Core functionality for generating DOCX resumes from YAML files.
"""

import os
from typing import Optional, Tuple, Dict, Any, List
from datetime import datetime
from docx import Document
from io import BytesIO

from docx.shared import RGBColor, Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.text.run import Run

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.document import Document as DocxDocument
from .converter import docx_to_pdf
from .validator import load_yaml
from .bullets import BULLETS

def count_pages(doc: DocxDocument) -> int:
    """Count the number of pages in a document.
    
    This method uses LibreOffice to get an accurate page count by converting to PDF temporarily.
    The PDF is deleted after counting pages.
    """
    import tempfile
    import os
    import subprocess
    
    # Create a temporary file for the PDF
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        temp_docx_path = temp_docx.name
        doc.save(temp_docx_path)
    
    try:
        # Convert to PDF using LibreOffice
        pdf_path = os.path.splitext(temp_docx_path)[0] + '.pdf'
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to',
            'pdf',
            temp_docx_path,
            '--outdir',
            os.path.dirname(temp_docx_path)
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        # Use pdfinfo to get page count
        result = subprocess.run(['pdfinfo', pdf_path], 
                              capture_output=True, 
                              text=True)
        for line in result.stdout.split('\n'):
            if line.startswith('Pages:'):
                page_count = int(line.split()[1])
                break
        else:
            page_count = 1  # fallback if pages not found
            
        return page_count
        
    finally:
        # Clean up temporary files
        try:
            os.unlink(temp_docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
        except Exception:
            pass  # ignore cleanup errors

def add_page_numbers(doc: DocxDocument) -> None:
    """Add page numbers in the format 'Page X of Y' to the footer."""
    section = doc.sections[0]
    footer = section.footer
    
    # Create footer paragraph with page numbers
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add "Page "
    paragraph.add_run("Page ")
    
    # Add page number field
    page_num = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    
    page_num._r.append(fld_char1)
    page_num._r.append(instr_text)
    page_num._r.append(fld_char2)
    
    # Add " of "
    paragraph.add_run(" of ")
    
    # Add total pages field
    num_pages = paragraph.add_run()
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'begin')
    
    instr_text2 = OxmlElement('w:instrText')
    instr_text2.set(qn('xml:space'), 'preserve')
    instr_text2.text = 'NUMPAGES'
    
    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')
    
    num_pages._r.append(fld_char3)
    num_pages._r.append(instr_text2)
    num_pages._r.append(fld_char4)

def setup_document() -> DocxDocument:
    """Create and configure a new document with default settings."""
    doc = Document()
    # Set margins for compact layout
    for section in doc.sections:
        section.top_margin = Inches(0.15)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

    # Font defaults
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(9)
    doc.styles['Heading 1'].font.size = Pt(12)
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(4)
    doc.styles['Heading 2'].font.size = Pt(10)
    doc.styles['Title'].paragraph_format.space_after = Pt(2)

    #breakpoint()

    if "MySectionStyle" not in doc.styles:
        new_style = doc.styles.add_style("MySectionStyle", WD_STYLE_TYPE.PARAGRAPH)
        base = doc.styles["Normal"]
        new_style.base_style = base
        new_style.paragraph_format.space_after = Pt(0)

    if "MyBulletStyle" not in doc.styles:
        new_style = doc.styles.add_style("MyBulletStyle", WD_STYLE_TYPE.PARAGRAPH)
        base = doc.styles["List Bullet"]
        new_style.base_style = base
        new_style.paragraph_format.space_after = Pt(10)

    if "CustomLabel" not in doc.styles:
        font = doc.styles.add_style("CustomLabel", WD_STYLE_TYPE.CHARACTER).font
        #font.name = 'Arial'
        #font.size = Pt(12)
        font.bold = True
        #font.color.rgb = RGBColor(0x36, 0x5F, 0x91)  # Blue
        font.color.rgb = doc.styles['Title'].font.color.rgb
        
    if "Hyperlink" not in doc.styles:
        hs = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
        hs.unhide_when_used = True
        hs.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        
    # Add page numbers to footer
    add_page_numbers(doc)
    
    return doc

def _is_year(date: str) -> bool:
    """Returns True if string is a 4-digit year, False otherwise."""
    return isinstance(date, str) and date.isdigit() and len(date) == 4

def get_job_dates(job: Dict[str, Any]) -> List[str]:
    """Extract start and end dates from a job entry."""
    dates = []
    if "startDate" in job:
        dates.append(job["startDate"])
    if "endDate" in job:
        dates.append(job["endDate"])
    return dates

def format_date_range(dates: List[str]) -> str:
    """Format a list of dates into a date range string."""
    formatted_dates = []
    for date in dates:
        if date.lower() == "present":
            formatted_dates.append("Present")
        elif _is_year(date):
            formatted_dates.append(date)
        else:
            formatted_dates.append(datetime.strptime(date, "%Y-%m-%d").strftime("%b %Y"))
    if len(formatted_dates) > 1 and formatted_dates[0] == formatted_dates[1]:
        formatted_dates = [formatted_dates[0]]
    return f" ({' - '.join(formatted_dates)})" if formatted_dates else ""

def add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = Run(OxmlElement('w:r'), paragraph)
    new_run.text = text
    new_run.style = "Hyperlink"
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_basics_section(doc: DocxDocument, basics: Dict[str, Any]) -> None:
    """Add the basic information section to the document."""
    if not basics:
        return
        
    # Name as heading
    heading = doc.add_heading(basics.get("name", ""), 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact information
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_parts = []
    if "email" in basics:
        add_hyperlink(contact, basics["email"], f"mailto:{basics['email']}")
    if "phone" in basics:
        contact.add_run(f" {BULLETS['MIDDLE_DOT']} " + basics["phone"] + f" {BULLETS['MIDDLE_DOT']} ")
    if "url" in basics:
        add_hyperlink(contact, basics["url"].replace("https://", ""), basics["url"])
    
    # Professional summary
    if "summary" in basics:
        doc.add_heading("Professional Summary", 1)
        doc.add_paragraph(basics["summary"])

def copy_doc(doc: DocxDocument) -> DocxDocument:
    """Create a copy of a document for page count testing."""
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Document(buffer)

def add_page_break_after_paragraph(paragraph):
    """Add a page break after the specified paragraph."""
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def _filter_ats_chars(text: str) -> str:
    """Remove characters that might confuse ATS systems."""
    return ''.join(c for c in text if c not in '[]{}()<>:|')

def ats_format_work_entry(job: Dict[str, Any]) -> Dict[str, Any]:
    """Format a work entry in ATS-friendly format."""
    modified_job = job.copy()  # Create a copy of the original job
    # Override specific fields for ATS format
    modified_job.update({
        "position": _filter_ats_chars(job.get("position", "")),
        "name": _filter_ats_chars(job.get("name", "")),
    })
    return modified_job

def add_work_entry(doc: DocxDocument, job: Dict[str, Any], style: str = "MyBulletStyle", ats_format: bool = False) -> DocxDocument:
    """Add a single work entry to the document."""
    if ats_format:
        # ATS-friendly format: title, company, and dates on separate lines
        ats_job = ats_format_work_entry(job)
        p = doc.add_paragraph(style="MySectionStyle")
        p.add_run(ats_job.get("position", "")).style = doc.styles["CustomLabel"]
        p.add_run("\n")
        p.add_run(ats_job.get("name", "")).bold = False
        p.add_run("\n")
        p.add_run(format_date_range(get_job_dates(job)))
    else:
        # Standard format: company | title | dates on one line
        p = doc.add_paragraph(style="MySectionStyle")
        p.add_run(job.get("name", "")).bold = False
        p.add_run(" | ")
        p.add_run(job.get("position", "")).style = doc.styles["CustomLabel"]
        p.add_run(" | ")
        p.add_run(format_date_range(get_job_dates(job)))
    
    if not ats_format and "summary" in job:
        doc.add_paragraph(job["summary"], style="MySectionStyle")
    
    last_p = p
    if "highlights" in job:
        for highlight in job["highlights"]:
            last_p = doc.add_paragraph(highlight, style=style)
    return last_p

def _is_page_break_needed(doc: DocxDocument, job: Dict[str, Any], ats_format: bool = False) -> bool:
    """Check if adding this job entry would cause a page break."""
    test_doc = copy_doc(doc)
    pages_before = count_pages(test_doc)
    add_work_entry(test_doc, job, ats_format=ats_format)
    pages_after = count_pages(test_doc)
    return pages_after > pages_before

def add_work_section(doc: DocxDocument, work_experience: List[Dict[str, Any]], ats_format: bool = False) -> None:
    """Add the work experience section to the document."""
    if not work_experience:
        return
        
    added_heading = False
    previous_paragraph = None
    
    for job in work_experience:
        if job.get("display", "") == "None":
            continue
        if ats_format and job.get("display", "") == "ATS-None":
            continue
        if not ats_format and job.get("display", "") == "ATS-Only":
            continue
            
        if not added_heading:
            doc.add_heading("Work Experience", 1)
            added_heading = True

        # If it would cause a page break, add one before the entry
        if not ats_format and previous_paragraph and _is_page_break_needed(doc, job, ats_format=ats_format):
            add_page_break_after_paragraph(previous_paragraph)
            print(f"Added page break before {job.get('position', '')} at {job.get('name', '')}")
        
        previous_paragraph = add_work_entry(doc, job, ats_format=ats_format)

def add_education_section(doc: DocxDocument, education: List[Dict[str, Any]]) -> None:
    """Add the education section to the document."""
    if not education:
        return
        
    added_heading = False
    for edu in education:
        if edu.get("display", "") == "None":
            continue
        if not added_heading:
            doc.add_heading("Education", 1)
            added_heading = True

        title = f"{edu.get('studyType', '')}"
        if title:
            title += f" in "
        title += f"{edu.get('area', '')}"
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(f" from {edu.get('institution', '')}")
        
        dates = []
        if "startDate" in edu:
            dates.append(edu["startDate"])
        if "endDate" in edu:
            dates.append(edu["endDate"])
        p.add_run(format_date_range(dates))

def add_skills_section(doc: DocxDocument, skills: List[Dict[str, Any]]) -> None:
    """Add the skills section to the document."""
    if not skills:
        return
        
    doc.add_heading("Skills", 1)
    for skill in skills:
        #p = doc.add_paragraph()
        p = doc.add_paragraph(style="MySectionStyle")
        if "name" in skill:
            p.add_run(skill["name"]).style = doc.styles["CustomLabel"]
        if "keywords" in skill:
            p.add_run(": " + f" {BULLETS['MIDDLE_DOT']} ".join(skill["keywords"]))

def generate_resume(
    yaml_file: str,
    output_file: str,
    convert_to_pdf: bool = False,
    ats_format: bool = False
) -> Tuple[bool, str, Optional[int]]:
    """
    Generate a DOCX resume from a YAML file following the JSON Resume schema.
    
    Args:
        yaml_file (str): Path to the input YAML file
        output_file (str): Path to the output DOCX file
        convert_to_pdf (bool): Whether to also create a PDF version
        ats_format (bool): Whether to use ATS-friendly formatting for work entries
    
    Returns:
        Tuple[bool, str, Optional[int]]: Success status, message, and page count (if successful)
    """
    try:
        # Load and validate YAML
        resume_data = load_yaml(yaml_file)
        
        # Create and setup document
        doc = setup_document()
        
        # Add each section
        add_basics_section(doc, resume_data.get("basics", {}))
        add_work_section(doc, resume_data.get("work", []), ats_format=ats_format)
        add_education_section(doc, resume_data.get("education", []))
        add_skills_section(doc, resume_data.get("skills", []))
        
        # Get page count before saving
        page_count = count_pages(doc)
        
        # Save the document
        doc.save(output_file)
        
        # Convert to PDF if requested
        if convert_to_pdf:
            pdf_success, pdf_message = docx_to_pdf(output_file)
            if not pdf_success:
                return False, f"DOCX created but PDF conversion failed: {pdf_message}", None
            return True, f"Successfully created {output_file} and PDF version ({page_count} pages)", page_count
        
        return True, f"Successfully created {output_file} ({page_count} pages)", page_count
    
    except Exception as e:
        return False, f"Error generating resume: {str(e)}", None 
