"""
Tests for the docx_generator.converter module.
"""

import os
import pytest
from docx import Document
from docx_generator.converter import docx_to_pdf

def test_nonexistent_file():
    """Test conversion with a nonexistent file."""
    success, message = docx_to_pdf("nonexistent.docx")
    assert not success
    assert "Input file 'nonexistent.docx' not found" == message

def test_invalid_extension(tmp_path):
    """Test conversion with an invalid file extension."""
    test_txt = tmp_path / "test.txt"
    test_txt.write_text("This is a test file")
    
    success, message = docx_to_pdf(str(test_txt))
    assert not success
    assert f"Input file '{test_txt}' is not a .docx file" == message

@pytest.mark.integration
def test_conversion(tmp_path):
    """
    Test actual DOCX to PDF conversion.
    
    This test requires LibreOffice to be installed and a test.docx file to exist.
    Mark as integration test since it depends on external software.
    """
    # Create a simple test DOCX file with actual content
    test_docx = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.save(str(test_docx))
    
    success, message = docx_to_pdf(str(test_docx))
    assert success, message
    assert os.path.exists(str(test_docx).replace(".docx", ".pdf")) 